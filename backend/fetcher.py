"""
获取学校列表和招生简章内容
"""

import httpx
import asyncio
import re
import io
import base64
import time
import tempfile
import subprocess
from charset_normalizer import from_bytes as detect_encoding
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse, quote_plus, parse_qs, unquote
from pypdf import PdfReader
import os
from config import PDF_DIR, REQUEST_DELAY_MS

try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None

try:
    import openpyxl
except Exception:
    openpyxl = None

try:
    import xlrd
except Exception:
    xlrd = None


def decode_bytes_safely(content: bytes, content_type: str = "") -> str:
    """尽量避免乱码地解码网页/文本字节。"""
    if not content:
        return ""

    # 先读 header 中的 charset
    m = re.search(r"charset=([a-zA-Z0-9_\-]+)", content_type or "", re.I)
    candidates = []
    if m:
        candidates.append(m.group(1).strip())

    # 自动探测
    try:
        best = detect_encoding(content).best()
        if best and best.encoding:
            candidates.append(best.encoding)
    except Exception:
        pass

    # 常见中文编码兜底
    candidates.extend(["utf-8", "gb18030", "gbk", "big5"])

    seen = set()
    for enc in candidates:
        key = (enc or "").lower()
        if not key or key in seen:
            continue
        seen.add(key)
        try:
            return content.decode(enc, errors="strict")
        except Exception:
            continue

    # 最后的宽松回退
    return content.decode("utf-8", errors="ignore")


# 确保PDF目录存在
os.makedirs(PDF_DIR, exist_ok=True)

HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
}


def normalize_fetch_url(raw_url: str) -> str:
    """清洗可能包含前缀文本/符号的URL。"""
    url = (raw_url or "").strip()
    if not url:
        return ""

    m = re.search(r"https?://[^\s\"'<>]+", url)
    if m:
        url = m.group(0)

    if url.startswith("//"):
        url = "https:" + url

    url = url.strip(" \t\r\n\"'<>[]()（）【】,;，；。")
    return url


def extract_front_cms_content_id(url: str, html: str = "") -> str:
    """从 front cms 页面 URL/HTML 中提取 contentId。"""
    parsed = urlparse(url or "")
    qs = parse_qs(parsed.query)
    for key in ["id", "contentId", "contentid", "articleId", "articleid"]:
        vals = qs.get(key) or []
        if vals and re.fullmatch(r"[0-9a-fA-F]{32}", vals[0] or ""):
            return vals[0]

    m = re.search(
        r"contentId\s*=\s*\$\.getAddrUrl\('id'\)\|\|'([0-9a-fA-F]{32})'", html or ""
    )
    if m:
        return m.group(1)

    m2 = re.search(r"[?&]id=([0-9a-fA-F]{32})", html or "")
    if m2:
        return m2.group(1)

    return ""


def html_fragment_to_text(fragment_html: str) -> str:
    soup = BeautifulSoup(fragment_html or "", "lxml")
    for tag in soup(["script", "style"]):
        tag.decompose()
    lines = [
        line.strip()
        for line in soup.get_text("\n", strip=True).split("\n")
        if line.strip()
    ]
    return "\n".join(lines)


def build_attachment_candidates(article: dict, base_url: str) -> list:
    links = []
    seen = set()
    for att in article.get("attachment") or []:
        if not isinstance(att, dict):
            continue

        raw_url = (
            att.get("url")
            or att.get("attachmentUrl")
            or att.get("attachmentPath")
            or att.get("path")
            or att.get("filePath")
            or ""
        )
        if not isinstance(raw_url, str) or not raw_url.strip():
            continue

        full = normalize_fetch_url(urljoin(base_url, raw_url.strip()))
        if not full or full in seen:
            continue
        seen.add(full)

        name = (
            att.get("name")
            or att.get("fileName")
            or att.get("attachmentName")
            or att.get("title")
            or full
        )
        links.append({"url": full, "text": str(name)})

    return links


async def fetch_front_cms_ajax_content(
    page_url: str,
    html: str,
    client: httpx.AsyncClient,
) -> dict:
    """针对 /static/front/* 模板页，调用其 ajax_article_view 获取正文。"""
    cid = extract_front_cms_content_id(page_url, html)
    if not cid:
        return {"success": False, "text": "", "pdf_links": [], "error": "缺少contentId"}

    parsed = urlparse(page_url)
    base = f"{parsed.scheme}://{parsed.netloc}/"
    article_api = urljoin(base, "f/newsCenter/ajax_article_view")
    token_api = urljoin(base, "f/ajax_get_csrfToken")

    try:
        # 先做一次预请求，部分站点会在响应中回传 jessionid
        pre = await client.get(
            f"{article_api}?contentId={cid}",
            headers={"X-Requested-With": "XMLHttpRequest"},
        )
        try:
            pre_json = pre.json()
            jsid = pre_json.get("jessionid")
            if isinstance(jsid, str) and jsid:
                client.cookies.set("JSESSIONID", jsid, domain=parsed.netloc, path="/")
                client.cookies.set("jessionid", jsid, domain=parsed.netloc, path="/")
        except Exception:
            pass

        ts = str(int(time.time() * 1000))
        tok_resp = await client.post(
            f"{token_api}?ts={ts}",
            data={"n": "3"},
            headers={
                "X-Requested-Time": ts,
                "X-Requested-With": "XMLHttpRequest",
                "Referer": page_url,
            },
        )

        token = ""
        try:
            tok_json = tok_resp.json()
            data = tok_json.get("data")
            if isinstance(data, str) and data:
                token = data.split(",")[0]
        except Exception:
            pass

        ts2 = str(int(time.time() * 1000))
        headers = {
            "X-Requested-Time": ts2,
            "X-Requested-With": "XMLHttpRequest",
            "Referer": page_url,
        }
        if token:
            headers["Csrf-Token"] = token

        ar = await client.post(
            f"{article_api}?ts={ts2}",
            data={"contentId": cid},
            headers=headers,
        )

        ar_json = ar.json()
        if ar_json.get("state") != 1 or not isinstance(ar_json.get("data"), dict):
            return {
                "success": False,
                "text": "",
                "pdf_links": [],
                "error": f"ajax_article_view state={ar_json.get('state')}",
            }

        article = (ar_json.get("data") or {}).get("article") or {}
        title = (article.get("title") or "").strip()
        release_date = article.get("releaseDate") or ""
        content_html = ((article.get("articleData") or {}).get("content") or "").strip()
        content_text = html_fragment_to_text(content_html)
        text_parts = []
        if title:
            text_parts.append(title)
        if release_date:
            text_parts.append(f"发布时间: {release_date}")
        if content_text:
            text_parts.append(content_text)

        text = "\n".join(text_parts).strip()
        atts = build_attachment_candidates(article, base)

        return {
            "success": bool(text),
            "text": text,
            "pdf_links": atts,
            "error": None if text else "ajax内容为空",
        }
    except Exception as e:
        return {
            "success": False,
            "text": "",
            "pdf_links": [],
            "error": str(e),
        }


async def fetch_school_list() -> list:
    """从官网API获取学校列表"""
    async with httpx.AsyncClient(
        headers=HEADERS, timeout=30, follow_redirects=True
    ) as client:
        resp = await client.get("https://www.gatzs.com.cn/xcbm/yxxx/getSsVOListJson")
        data = resp.json()

        schools = []
        for region in data.get("data", []):
            area_name = region.get("areaName", "")
            for school in region.get("schools", []):
                schools.append(
                    {
                        "id": school.get("id"),
                        "code": school.get("code"),
                        "name": school.get("name"),
                        "type": school.get("yxtx", ""),
                        "area": area_name,
                        "zsjz_url": school.get("zsjz"),  # 招生简章URL
                    }
                )
        return schools


async def fetch_page_content(url: str, client: httpx.AsyncClient) -> dict:
    """获取页面内容，自动处理HTML和PDF"""
    clean_url = normalize_fetch_url(url)
    result = {
        "url": clean_url or url,
        "success": False,
        "content_type": None,
        "text": "",
        "html": None,
        "image_links": [],
        "error": None,
    }

    if not clean_url:
        result["error"] = "无URL"
        return result

    try:

        async def _get_with_retry(u: str) -> httpx.Response:
            last = None
            for attempt in range(2):
                try:
                    return await client.get(u, follow_redirects=True)
                except Exception as e:
                    last = e
                    emsg = str(e)
                    retryable = any(
                        k in emsg
                        for k in [
                            "Server disconnected without sending a response",
                            "All connection attempts failed",
                            "ConnectTimeout",
                            "ReadTimeout",
                            "timed out",
                            "timeout",
                            "RemoteProtocolError",
                        ]
                    )
                    if attempt == 0 and retryable:
                        await asyncio.sleep(0.8)
                        continue
                    raise

            raise last

        try:
            resp = await _get_with_retry(clean_url)
        except Exception as e:
            # 某些老站 HTTPS 握手失败，回退到 HTTP 再试一次
            emsg = str(e)
            can_http_retry = clean_url.startswith("https://") and any(
                k in emsg
                for k in [
                    "UNSAFE_LEGACY_RENEGOTIATION_DISABLED",
                    "DH_KEY_TOO_SMALL",
                    "SSL",
                    "Server disconnected without sending a response",
                ]
            )
            if not can_http_retry:
                raise

            fallback_url = "http://" + clean_url[len("https://") :]
            resp = await _get_with_retry(fallback_url)
            clean_url = fallback_url
            result["url"] = fallback_url

        content_type = resp.headers.get("content-type", "").lower()
        result["content_type"] = content_type
        lower_url = clean_url.lower()

        # PDF处理
        if "application/pdf" in content_type or lower_url.endswith(".pdf"):
            result["text"] = await extract_pdf_text(resp.content, clean_url)
            result["success"] = bool((result["text"] or "").strip())

        # DOCX 处理
        elif (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            in content_type
            or lower_url.endswith(".docx")
        ):
            result["text"] = extract_docx_text(resp.content)
            result["success"] = bool((result["text"] or "").strip())

        # DOC 处理
        elif "application/msword" in content_type or lower_url.endswith(".doc"):
            result["text"] = extract_doc_text(resp.content)
            result["success"] = bool((result["text"] or "").strip())

        # XLSX 处理
        elif (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            in content_type
            or lower_url.endswith(".xlsx")
        ):
            result["text"] = extract_xlsx_text(resp.content)
            result["success"] = bool((result["text"] or "").strip())

        # XLS 处理
        elif "application/vnd.ms-excel" in content_type or lower_url.endswith(".xls"):
            result["text"] = extract_xls_text(resp.content)
            result["success"] = bool((result["text"] or "").strip())

        # HTML处理
        elif "text/html" in content_type:
            html = decode_bytes_safely(resp.content, content_type)
            result["html"] = html
            result["text"] = extract_html_text(html, clean_url)
            result["success"] = True

            # 针对 static/front 模板页，尝试通过 ajax_article_view 获取真实正文
            short_text = len((result.get("text") or "").strip()) < 80
            front_cms_hint = (
                "/static/front/" in lower_url
                or "ajax_article_view" in html
                or "$.url('f/newsCenter/ajax_article_view')" in html
            )
            if front_cms_hint and short_text:
                ajax_result = await fetch_front_cms_ajax_content(
                    clean_url, html, client
                )
                if ajax_result.get("success") and len(
                    (ajax_result.get("text") or "").strip()
                ) > len((result.get("text") or "").strip()):
                    result["text"] = ajax_result.get("text") or result["text"]
                    result["success"] = True

                ajax_links = ajax_result.get("pdf_links") or []
                if ajax_links:
                    result["pdf_links"] = ajax_links

            # 检查页面中是否有PDF链接（招生简章常见模式）
            pdf_links = find_pdf_links(html, clean_url)
            if pdf_links:
                existing = result.get("pdf_links") or []
                merged = []
                seen = set()
                for item in existing + pdf_links:
                    u = (item.get("url") or "").strip()
                    if not u or u in seen:
                        continue
                    seen.add(u)
                    merged.append(item)
                result["pdf_links"] = merged

            img_links = find_image_links(html, clean_url)
            if img_links:
                result["image_links"] = img_links

        # 内容类型不明确，但 URL 明确是附件
        elif any(
            lower_url.endswith(ext)
            for ext in [".pdf", ".doc", ".docx", ".xls", ".xlsx"]
        ):
            if lower_url.endswith(".pdf"):
                result["text"] = await extract_pdf_text(resp.content, clean_url)
            elif lower_url.endswith(".docx"):
                result["text"] = extract_docx_text(resp.content)
            elif lower_url.endswith(".doc"):
                result["text"] = extract_doc_text(resp.content)
            elif lower_url.endswith(".xlsx"):
                result["text"] = extract_xlsx_text(resp.content)
            elif lower_url.endswith(".xls"):
                result["text"] = extract_xls_text(resp.content)
            result["success"] = bool((result["text"] or "").strip())

        else:
            result["text"] = decode_bytes_safely(resp.content, content_type)[:10000]
            result["success"] = True

        if not result["success"] and not result.get("error"):
            result["error"] = "内容解析失败或为空"

    except Exception as e:
        result["error"] = str(e)

    return result


def extract_docx_text(content: bytes) -> str:
    if DocxDocument is None:
        return "DOCX解析失败: python-docx 未安装"
    try:
        doc = DocxDocument(io.BytesIO(content))
        parts = []

        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)

        # 读取表格内容
        for tb in doc.tables:
            for row in tb.rows:
                cells = []
                for c in row.cells:
                    v = (c.text or "").strip()
                    if v:
                        cells.append(v)
                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts)
    except Exception as e:
        return f"DOCX解析失败: {e}"


def extract_doc_text(content: bytes) -> str:
    """使用 antiword 提取 .doc 文本。"""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        try:
            proc = subprocess.run(
                ["antiword", tmp_path],
                capture_output=True,
                timeout=40,
                check=False,
            )
            stdout = proc.stdout or b""
            stderr = proc.stderr or b""

            out_text = ""
            if stdout:
                # antiword 中文输出常见 gb18030/utf-8
                for enc in ["utf-8", "gb18030", "gbk", "big5"]:
                    try:
                        out_text = stdout.decode(enc)
                        break
                    except Exception:
                        continue
                if not out_text:
                    out_text = stdout.decode("utf-8", errors="ignore")

            if proc.returncode == 0 and out_text.strip():
                return out_text

            err = ""
            if stderr:
                err = stderr.decode("utf-8", errors="ignore").strip()
            if err:
                return f"DOC解析失败: {err}"
            return "DOC解析失败: antiword返回空内容"
        finally:
            try:
                os.remove(tmp_path)
            except Exception:
                pass
    except Exception as e:
        return f"DOC解析失败: {e}"


def extract_xlsx_text(content: bytes) -> str:
    if openpyxl is None:
        return "XLSX解析失败: openpyxl 未安装"

    try:
        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        lines = []
        for ws in wb.worksheets[:5]:
            lines.append(f"[Sheet] {ws.title}")
            row_count = 0
            for row in ws.iter_rows(values_only=True):
                vals = [str(v).strip() for v in row if v is not None and str(v).strip()]
                if vals:
                    lines.append(" | ".join(vals))
                row_count += 1
                if row_count >= 500:
                    break
        return "\n".join(lines)
    except Exception as e:
        return f"XLSX解析失败: {e}"


def extract_xls_text(content: bytes) -> str:
    if xlrd is None:
        return "XLS解析失败: xlrd 未安装"
    try:
        book = xlrd.open_workbook(file_contents=content)
        lines = []
        for si in range(min(book.nsheets, 5)):
            sh = book.sheet_by_index(si)
            lines.append(f"[Sheet] {sh.name}")
            for ri in range(min(sh.nrows, 500)):
                vals = []
                for ci in range(sh.ncols):
                    v = sh.cell_value(ri, ci)
                    if v is None:
                        continue
                    sv = str(v).strip()
                    if sv:
                        vals.append(sv)
                if vals:
                    lines.append(" | ".join(vals))
        return "\n".join(lines)
    except Exception as e:
        return f"XLS解析失败: {e}"


def discover_candidate_links(html: str, base_url: str, school_name: str = "") -> list:
    """从页面中发现更可能包含 2026 台湾学测简章细节的链接。"""
    soup = BeautifulSoup(html, "lxml")
    links = []
    seen = set()

    school_key = school_name[:4] if school_name else ""

    for a in soup.find_all("a", href=True):
        href = (a.get("href") or "").strip()
        if not href:
            continue

        full_url = urljoin(base_url, href)
        if full_url in seen:
            continue
        seen.add(full_url)

        text = (a.get_text(" ", strip=True) or "").strip()
        combo = f"{text} {href}"

        score = 0
        if "2026" in combo or "26年" in combo:
            score += 4
        if any(k in combo for k in ["台湾", "台灣", "港澳台"]):
            score += 4
        if any(k in combo for k in ["学测", "學測"]):
            score += 4
        if any(
            k in combo
            for k in [
                "学科能力测试",
                "学科能力测验",
                "學科能力測驗",
                "學科能力測試",
            ]
        ):
            score += 3
        if any(k in combo for k in ["简章", "簡章", "招生", "招生简章"]):
            score += 3
        if school_name and (
            school_name in combo or (school_key and school_key in combo)
        ):
            score += 2

        lower = href.lower()
        is_attachment = any(
            ext in lower
            for ext in [
                ".pdf",
                ".doc",
                ".docx",
                ".xls",
                ".xlsx",
                ".png",
                ".jpg",
                ".jpeg",
            ]
        )
        if any(k in lower for k in ["download", "attachment", "downattach", "fileid="]):
            is_attachment = True
        if is_attachment:
            score += 2

        if any(k in text for k in ["附件", "下载", "點擊下載", "点击下载"]):
            score += 2

        # 对明显无关链接降权
        if any(
            k in lower for k in ["javascript:", "mailto:", "weibo.com", "zhihu.com"]
        ):
            score -= 5

        if score >= 6:
            links.append(
                {
                    "url": full_url,
                    "text": text,
                    "score": score,
                    "is_attachment": is_attachment,
                }
            )

    links.sort(key=lambda x: x["score"], reverse=True)
    return links


async def fetch_related_content(
    base_url: str,
    html: str,
    client: httpx.AsyncClient,
    school_name: str = "",
    max_links: int = 4,
) -> dict:
    """从入口页中继续追踪候选详情/附件链接并抓取文本。"""
    candidates = discover_candidate_links(html, base_url, school_name)
    if not candidates:
        return {"success": False, "text": "", "links": [], "image_links": []}

    chunks = []
    used_links = []
    image_links = []

    def merge_images(links: list):
        for u in links or []:
            if u and u not in image_links:
                image_links.append(u)

    for item in candidates[:max_links]:
        try:
            await asyncio.sleep(REQUEST_DELAY_MS / 1000)
            sub = await fetch_page_content(item["url"], client)
            if not sub.get("success"):
                continue

            merge_images(sub.get("image_links") or [])
            t = (sub.get("text") or "").strip()
            if len(t) >= 60:
                used_links.append(item)
                chunks.append(f"\n\n--- 子链接内容: {item['url']} ---\n{t}")

            # 二跳页面如果有 PDF 链接，再补抓 1 份
            for pdf in (sub.get("pdf_links") or [])[:1]:
                pdf_text = pdf.get("text") or ""
                pdf_url = (pdf.get("url") or "").lower()
                if any(
                    k in pdf_text for k in ["2026", "台湾", "学测", "简章", "附件"]
                ) or any(
                    k in pdf_url
                    for k in [
                        ".pdf",
                        ".doc",
                        ".docx",
                        ".xls",
                        ".xlsx",
                        "download",
                        "attachment",
                    ]
                ):
                    await asyncio.sleep(REQUEST_DELAY_MS / 1000)
                    pdf_sub = await fetch_page_content(pdf["url"], client)
                    merge_images(pdf_sub.get("image_links") or [])
                    if pdf_sub.get("success") and (pdf_sub.get("text") or "").strip():
                        chunks.append(
                            f"\n\n--- 子链接PDF内容: {pdf['url']} ---\n{pdf_sub.get('text', '')}"
                        )
                    break

            # 若子链接本身是列表页，继续追踪一层详情链接
            if sub.get("html"):
                second_level = discover_candidate_links(
                    sub.get("html") or "", item["url"], school_name
                )
                for child in second_level[:2]:
                    if child["url"] == item["url"]:
                        continue
                    await asyncio.sleep(REQUEST_DELAY_MS / 1000)
                    child_result = await fetch_page_content(child["url"], client)
                    if not child_result.get("success"):
                        continue
                    merge_images(child_result.get("image_links") or [])
                    child_text = (child_result.get("text") or "").strip()
                    if len(child_text) >= 80:
                        used_links.append(child)
                        chunks.append(
                            f"\n\n--- 二跳链接内容: {child['url']} ---\n{child_text}"
                        )

                    # 二跳页面附件再补抓 1 份
                    for child_att in (child_result.get("pdf_links") or [])[:1]:
                        child_text_flag = child_att.get("text") or ""
                        child_url_flag = (child_att.get("url") or "").lower()
                        if any(
                            k in child_text_flag
                            for k in ["2026", "台湾", "学测", "简章", "附件"]
                        ) or any(
                            k in child_url_flag
                            for k in [
                                ".pdf",
                                ".doc",
                                ".docx",
                                ".xls",
                                ".xlsx",
                                "download",
                                "attachment",
                            ]
                        ):
                            await asyncio.sleep(REQUEST_DELAY_MS / 1000)
                            att_result = await fetch_page_content(
                                child_att["url"], client
                            )
                            merge_images(att_result.get("image_links") or [])
                            if (
                                att_result.get("success")
                                and (att_result.get("text") or "").strip()
                            ):
                                chunks.append(
                                    f"\n\n--- 二跳链接附件内容: {child_att['url']} ---\n{att_result.get('text', '')}"
                                )
                            break
        except Exception:
            continue

    if not chunks:
        return {
            "success": False,
            "text": "",
            "links": candidates[:max_links],
            "image_links": image_links,
        }

    return {
        "success": True,
        "text": "\n".join(chunks),
        "links": used_links,
        "image_links": image_links,
    }


def extract_html_text(html: str, base_url: str) -> str:
    """从HTML提取正文文本"""
    soup = BeautifulSoup(html, "lxml")

    # 移除脚本和样式
    for tag in soup(["script", "style", "nav", "header", "footer", "aside"]):
        tag.decompose()

    # 尝试找到正文区域
    main_content = None
    for selector in [
        "article",
        "main",
        ".content",
        ".article",
        "#content",
        ".main-content",
        ".news-content",
        ".detail",
    ]:
        main_content = soup.select_one(selector)
        if main_content:
            break

    if main_content:
        text = main_content.get_text(separator="\n", strip=True)
    else:
        text = soup.get_text(separator="\n", strip=True)

    # 清理多余空行
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    return "\n".join(lines)


def find_pdf_links(html: str, base_url: str) -> list:
    """查找页面中的附件/简章链接（保留字段名 pdf_links 兼容旧逻辑）。"""
    soup = BeautifulSoup(html, "lxml")
    pdf_links = []
    seen = set()

    for a in soup.find_all("a", href=True):
        href = (a["href"] or "").strip()
        text = a.get_text(strip=True)
        if not href:
            continue

        lower = href.lower()
        is_attachment = any(
            ext in lower for ext in [".pdf", ".doc", ".docx", ".xls", ".xlsx"]
        ) or any(
            k in lower for k in ["download", "attachment", "downattach", "fileid="]
        )

        # 附件或关键词链接均纳入候选
        if is_attachment or any(
            kw in (text + href) for kw in ["简章", "招生", "台湾", "学测", "港澳台"]
        ):
            full_url = urljoin(base_url, href)
            if full_url in seen:
                continue
            seen.add(full_url)
            pdf_links.append({"url": full_url, "text": text})

    return pdf_links


def find_image_links(html: str, base_url: str) -> list:
    """查找页面中的图片链接（优先可能是表格截图的资源）。"""
    soup = BeautifulSoup(html, "lxml")
    links = []
    seen = set()
    for img in soup.find_all("img", src=True):
        src = (img.get("src") or "").strip()
        if not src:
            continue
        full = urljoin(base_url, src)
        if full in seen:
            continue
        seen.add(full)

        lower = full.lower()
        if any(ext in lower for ext in [".png", ".jpg", ".jpeg", ".webp"]):
            if any(
                k in lower
                for k in [
                    "logo",
                    "banner",
                    "footer",
                    "icon",
                    "menu",
                    "close",
                    "wx",
                    "qrcode",
                    "zscode",
                ]
            ):
                continue
            links.append(full)

    return links[:8]


async def extract_pdf_text(content: bytes, url: str) -> str:
    """从PDF提取文本"""
    # 保存PDF以备查看
    filename = urlparse(url).path.split("/")[-1] or "document.pdf"
    if not filename.endswith(".pdf"):
        filename += ".pdf"
    filepath = os.path.join(PDF_DIR, filename)

    with open(filepath, "wb") as f:
        f.write(content)

    # 提取文本
    text_parts = []
    try:
        reader = PdfReader(io.BytesIO(content))
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
    except Exception as e:
        return f"PDF解析失败: {e}"

    return "\n".join(text_parts)


def unwrap_bing_redirect(url: str) -> str:
    """尽量将 Bing 跳转链接还原为真实目标 URL。"""
    try:
        parsed = urlparse(url)
        if "bing.com" not in parsed.netloc or "/ck/" not in parsed.path:
            return url

        qs = parse_qs(parsed.query)
        for key in ("u", "r"):
            values = qs.get(key)
            if not values:
                continue
            raw = values[0]

            if raw.startswith(("http://", "https://")):
                return raw

            candidate = raw
            if raw.startswith("a1"):
                candidate = raw[2:]
            elif raw.startswith("a"):
                candidate = raw[1:]

            # 常见情况：Bing 将 URL 做了 base64-url 编码
            try:
                padded = candidate + ("=" * (-len(candidate) % 4))
                decoded = base64.urlsafe_b64decode(padded.encode("utf-8")).decode(
                    "utf-8", errors="ignore"
                )
                if decoded.startswith(("http://", "https://")):
                    return decoded
            except Exception:
                pass

            # 兜底：尝试 URL decode
            decoded_url = unquote(raw)
            if decoded_url.startswith(("http://", "https://")):
                return decoded_url
    except Exception:
        pass

    return url


def get_domain_root(host: str) -> str:
    """提取域名主干，用于同组织域名匹配。"""
    host = (host or "").split(":")[0].strip().lower().strip(".")
    if not host:
        return ""

    parts = [p for p in host.split(".") if p]
    if len(parts) < 2:
        return host

    last2 = ".".join(parts[-2:])
    cn_suffix = {"com.cn", "edu.cn", "gov.cn", "org.cn", "net.cn", "ac.cn"}
    if len(parts) >= 3 and last2 in cn_suffix:
        return ".".join(parts[-3:])

    return last2


def is_blocked_domain(host: str) -> bool:
    host = (host or "").lower()
    blocked = [
        "zhihu.com",
        "zhidao.baidu.com",
        "jingyan.baidu.com",
        "baike.baidu.com",
        "wikipedia.org",
        "reddit.com",
        "microsoft.com",
        "google.com",
        "techcommunity.microsoft.com",
        "csdn.net",
        "weibo.com",
        "xiaohongshu.com",
        "bilibili.com",
        "39.net",
        "commentcamarche.net",
        "policyx.com",
        "kabu-sokuhou.com",
        "lahoratime.com",
        "ncert.nic.in",
        "tinhte.vn",
    ]
    return any(d in host for d in blocked)


async def search_fallback(
    school_name: str,
    client: httpx.AsyncClient,
    preferred_url: str = "",
) -> dict:
    """搜索引擎fallback - 使用Bing搜索，优先学校官方域名。"""
    preferred_host = urlparse(preferred_url or "").netloc.lower().split(":")[0]
    preferred_root = get_domain_root(preferred_host)

    queries = []
    if preferred_host:
        host_no_www = (
            preferred_host[4:] if preferred_host.startswith("www.") else preferred_host
        )
        queries.append(f"site:{host_no_www} {school_name} 2026 台湾 学测 简章")
        if preferred_root and preferred_root not in host_no_www:
            queries.append(f"site:{preferred_root} {school_name} 2026 台湾 学测 简章")

    queries.append(f"{school_name} 2026年招收台湾地区高中毕业生简章")
    queries.append(f"{school_name} 2026 学科能力测验 招收 台湾 高中 毕业生 简章")
    queries.append(
        f"{school_name} 2026 依据台湾地区大学入学考试 学科能力测试 成绩 招收 台湾 高中 毕业生 招生简章"
    )

    result = {
        "url": "",
        "success": False,
        "text": "",
        "error": None,
        "source": "search",
    }

    try:
        best_candidate = None
        seen_urls = set()
        school_key_short = school_name[:4] if len(school_name) >= 4 else school_name
        all_search_results = []

        for q in queries:
            search_url = f"https://www.bing.com/search?q={quote_plus(q)}"
            result["url"] = search_url

            resp = await client.get(search_url)
            soup = BeautifulSoup(resp.text, "lxml")

            # 提取搜索结果
            query_results = []
            for item in soup.select(".b_algo")[:8]:
                title_el = item.select_one("h2 a")
                snippet_el = item.select_one(".b_caption p")
                if not title_el:
                    continue
                href = title_el.get("href", "")
                query_results.append(
                    {
                        "title": title_el.get_text(strip=True),
                        "url": href,
                        "resolved_url": unwrap_bing_redirect(href),
                        "snippet": snippet_el.get_text(strip=True)
                        if snippet_el
                        else "",
                    }
                )

            if not query_results:
                continue

            all_search_results.extend(query_results)

            for item in query_results:
                candidate_url = item.get("resolved_url") or item.get("url")
                if not candidate_url or candidate_url in seen_urls:
                    continue
                seen_urls.add(candidate_url)

                parsed_candidate = urlparse(candidate_url)
                host = parsed_candidate.netloc.lower().split(":")[0]
                root = get_domain_root(host)
                same_org = bool(preferred_root and root and root == preferred_root)

                if is_blocked_domain(host):
                    continue

                title = item.get("title", "")
                snippet = item.get("snippet", "")
                combined_text = f"{title} {snippet}"
                name_hit = (
                    school_name in combined_text or school_key_short in combined_text
                )

                lower_url = candidate_url.lower()
                is_edu_domain = ".edu.cn" in host or host.endswith(".edu")
                is_gov_domain = ".gov.cn" in host
                is_admission_like = any(
                    kw in (host + " " + lower_url)
                    for kw in [
                        "zsb",
                        "admission",
                        "zs.",
                        "bkzs",
                        "recruit",
                        "zhaosheng",
                        "gat",
                    ]
                )
                preferred_source = (
                    same_org or is_edu_domain or is_gov_domain or is_admission_like
                )

                # 有官方域名时尽量收敛在同组织或高校/招生域；没有则至少命中学校名或高校/招生域
                if preferred_root:
                    if not (same_org or is_edu_domain or is_admission_like or name_hit):
                        continue
                elif not (preferred_source or name_hit):
                    continue

                await asyncio.sleep(REQUEST_DELAY_MS / 1000)
                page_result = await fetch_page_content(candidate_url, client)
                if not page_result.get("success"):
                    continue

                text = (page_result.get("text") or "").strip()
                text_len = len(text)
                if text_len == 0:
                    continue

                content_name_hit = school_name in text or school_key_short in text
                keyword_hits = sum(
                    1
                    for kw in [
                        "2026",
                        "台湾",
                        "台灣",
                        "学测",
                        "學測",
                        "学科能力测试",
                        "学科能力测验",
                        "學科能力測驗",
                        "學科能力測試",
                        "招生",
                        "简章",
                        "簡章",
                        "港澳台",
                    ]
                    if kw in text
                )

                if keyword_hits == 0 and not (same_org and content_name_hit):
                    continue
                if not (content_name_hit or same_org) and keyword_hits < 2:
                    continue

                domain_bonus = 0
                if same_org:
                    domain_bonus += 1800
                if is_edu_domain:
                    domain_bonus += 1000
                if is_gov_domain:
                    domain_bonus += 600
                if is_admission_like:
                    domain_bonus += 700
                if name_hit:
                    domain_bonus += 700
                if content_name_hit:
                    domain_bonus += 900

                score = text_len + keyword_hits * 650 + domain_bonus
                rank = (1 if same_org else 0, 1 if preferred_source else 0, score)

                if best_candidate is None or rank > best_candidate["rank"]:
                    best_candidate = {
                        "rank": rank,
                        "score": score,
                        "text": text,
                        "final_url": candidate_url,
                        "text_len": text_len,
                        "keyword_hits": keyword_hits,
                        "image_links": page_result.get("image_links", []),
                    }

                # 同组织域名命中且关键词充足时可提前收敛
                if same_org and text_len >= 180 and keyword_hits >= 2:
                    result["success"] = True
                    result["text"] = text
                    result["final_url"] = candidate_url
                    result["text_len"] = text_len
                    result["keyword_hits"] = keyword_hits
                    result["image_links"] = page_result.get("image_links", [])
                    result["search_results"] = all_search_results
                    return result

        if all_search_results:
            result["search_results"] = all_search_results

        if best_candidate:
            result["success"] = True
            result["text"] = best_candidate["text"]
            result["final_url"] = best_candidate["final_url"]
            result["text_len"] = best_candidate["text_len"]
            result["keyword_hits"] = best_candidate["keyword_hits"]
            result["image_links"] = best_candidate.get("image_links", [])

    except Exception as e:
        result["error"] = str(e)

    return result
